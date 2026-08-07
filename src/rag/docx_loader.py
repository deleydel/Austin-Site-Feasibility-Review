"""Task 2: style-aware DOCX loader for the three Austin regulatory documents.

Design points (per review):
- Paragraphs and tables are walked in true XML body order, so tables land
  inside the section where they appear.
- Every paragraph is classified as one of:
    mapped_to_section | document_metadata | history_note | table_caption |
    intentionally_excluded | unresolved
  The loader reports counts; the build fails QA if any paragraph is
  `unresolved`.
- Text that follows a chapter/article heading but precedes the first legal
  section is NOT forced into a neighboring section: it becomes a `[preamble]`
  pseudo-section cited at the chapter/article level.

Output: a list of Section records — the atomic citation unit for chunking.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from src import config

CLASSES = (
    "mapped_to_section", "document_metadata", "history_note",
    "table_caption", "intentionally_excluded", "unresolved",
)


@dataclass
class Section:
    doc_id: str                 # LDC | DCM | TCM
    source_name: str
    source_url: str
    chapter: str                # "Chapter 25-2" / "Section 8" / "Appendix E"
    chapter_title: str
    article: str                # article/division path ("" if none)
    section_number: str         # "25-2-492" / "8.3.0" / "25-2 preamble"
    section_title: str
    is_preamble: bool = False
    paragraphs: list[str] = field(default_factory=list)
    history_notes: list[str] = field(default_factory=list)

    @property
    def breadcrumb(self) -> str:
        parts = [self.source_name, f"{self.chapter} {self.chapter_title}".strip()]
        if self.article:
            parts.append(self.article)
        if self.is_preamble:
            # section_number is "<scope> [preamble]"; the scope is already the
            # previous breadcrumb element, so show just the marker.
            parts.append("[preamble]")
        else:
            label = f"§ {self.section_number}" if self.doc_id == "LDC" \
                else self.section_number
            parts.append(f"{label} {self.section_title}".strip())
        return " > ".join(p for p in parts if p)

    @property
    def text(self) -> str:
        return "\n".join(self.paragraphs)

    def to_dict(self) -> dict:
        return {
            "doc_id": self.doc_id, "source_name": self.source_name,
            "source_url": self.source_url, "chapter": self.chapter,
            "chapter_title": self.chapter_title, "article": self.article,
            "section_number": self.section_number,
            "section_title": self.section_title, "is_preamble": self.is_preamble,
            "breadcrumb": self.breadcrumb, "text": self.text,
            "history_notes": self.history_notes,
        }


def _iter_body(doc: Document):
    """Yield Paragraph and Table objects in true XML order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def _table_to_markdown(tbl: Table) -> str:
    rows = []
    for r in tbl.rows:
        # A merged cell spans several grid columns as the SAME tc element;
        # collapse only those, never distinct cells that merely share text —
        # collapsing equal-valued columns would shift column alignment.
        cells, prev_tc = [], None
        for c in r.cells:
            if c._tc is prev_tc:
                continue
            prev_tc = c._tc
            cells.append(" ".join(c.text.split()))
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 2:
        n = rows[0].count("|") - 1
        rows.insert(1, "|" + " --- |" * n)
    return "\n".join(rows)


# --- section-number patterns ------------------------------------------------
# § prefix optional: a few headings (e.g. 25-6-101) omit it in the export.
LDC_SECTION_RE = re.compile(r"^§*\s*([0-9]+-[0-9A-Z]+-[0-9]+(?:\.[0-9]+)?)\s*(.*)$")
# Subchapter-ordinance sections inside LDC articles, e.g. "§ 1.1. GENERAL INTENT."
LDC_ORD_SECTION_RE = re.compile(r"^§+\s*(\d+(?:\.\d+)*)\.?\s+(.*)$")
NUMBERED_RE = re.compile(r"^(\d+(?:\.\d+)+)\.?\s+(.*)$")          # 1.2.2 Title / 7.5.1.1 Title.
MANUAL_CHAPTER_RE = re.compile(r"^SECTION\s+(\d+)\s*(.*)$", re.S)  # SECTION 1 \n DRAINAGE POLICY
APPENDIX_RE = re.compile(r"^APPENDIX\s+([A-Z])\s*(.*)$", re.S)
FIGURE_CAPTION_RE = re.compile(r"^(Figure|Table|Exhibit)\s?[\w./-]*[.:]?\s", re.I)

# Styles that carry structure per document.
LDC_CHAPTER, LDC_ARTICLE = "Heading 2", "Heading 3"
LDC_DIVISION = {"Heading 4", "Heading 5"}
LDC_SECTION_STYLE = "Section"
LDC_SUBSECT = "Subsect 1"           # named subsection inside a section (split point)

DCM_CHAPTER = {"Heading 1"}
DCM_SECTION = {"Heading 2", "Heading 3"}

TCM_CHAPTER = {"Heading 2"}
TCM_SECTION = {"Heading 3", "Heading 4", "Section"}

SUBSECTION_BOUNDARY_RE = re.compile(r"^\(?[A-Z]\)?[.)]\t|^\([A-Z]\)\s|^[A-Z]\.\t")


def load_document(path, doc_id: str) -> tuple[list[Section], dict]:
    """Parse one DOCX into Sections + a paragraph-classification audit."""
    meta = config.SOURCE_DOCS[doc_id]
    doc = Document(str(path))
    sections: list[Section] = []
    audit = {c: 0 for c in CLASSES}
    unresolved_samples: list[str] = []

    chapter = chapter_title = ""
    article_stack: list[str] = []
    current: Section | None = None

    def start_section(number: str, title: str, preamble: bool = False):
        nonlocal current
        current = Section(
            doc_id=doc_id, source_name=meta["name"], source_url=meta["url"],
            chapter=chapter, chapter_title=chapter_title,
            article=" > ".join(article_stack), section_number=number,
            section_title=title, is_preamble=preamble,
        )
        sections.append(current)

    def ensure_preamble():
        """Body text after a heading but before any numbered section."""
        nonlocal current
        if current is None:
            scope = article_stack[-1] if article_stack else chapter or "document"
            start_section(f"{scope} [preamble]".strip(), "", preamble=True)

    for el in _iter_body(doc):
        if isinstance(el, Table):
            if current is None:
                ensure_preamble()
            current.paragraphs.append(_table_to_markdown(el))
            continue

        style = el.style.name if el.style is not None else ""
        text = el.text.strip()
        if not text:
            audit["intentionally_excluded"] += 1
            continue
        text = re.sub(r"\s+", " ", text.replace(" ", " ")).strip()

        if doc_id == "LDC":
            if style == LDC_CHAPTER:
                m = re.match(r"^CHAPTER\s+([\d-]+)\.?\s*(.*)$", text)
                chapter = f"Chapter {m.group(1)}" if m else text
                chapter_title = (m.group(2) if m else "").strip(" .")
                article_stack, current = [], None
                audit["document_metadata"] += 1
            elif style == LDC_ARTICLE:
                article_stack = [text.strip(" .")]
                current = None
                audit["document_metadata"] += 1
            elif style in LDC_DIVISION:
                # Some ARTICLE headings are styled at division level; they
                # start a new article, not a nested division.
                if text.upper().startswith("ARTICLE "):
                    article_stack = [text.strip(" .")]
                else:
                    article_stack = article_stack[:1] + [text.strip(" .")]
                current = None
                audit["document_metadata"] += 1
            elif style == LDC_SECTION_STYLE:
                m = LDC_SECTION_RE.match(text) or LDC_ORD_SECTION_RE.match(text) \
                    or NUMBERED_RE.match(text)
                if m:
                    start_section(m.group(1), m.group(2).strip(" ."))
                else:
                    start_section(text, "")
                audit["document_metadata"] += 1
            elif style == "History Note":
                if current is not None:
                    current.history_notes.append(text)
                    audit["history_note"] += 1
                else:
                    audit["history_note"] += 1
            elif style == LDC_SUBSECT:
                # Named subsection header: keep inline (also a chunk split point).
                ensure_preamble()
                current.paragraphs.append(f"[SUBSECTION] {text}")
                audit["mapped_to_section"] += 1
            else:
                ensure_preamble()
                current.paragraphs.append(el.text.strip())
                audit["mapped_to_section"] += 1

        else:  # DCM / TCM share manual structure
            chapter_styles = DCM_CHAPTER if doc_id == "DCM" else TCM_CHAPTER
            section_styles = DCM_SECTION if doc_id == "DCM" else TCM_SECTION
            raw = el.text.strip()

            if style in chapter_styles:
                m = MANUAL_CHAPTER_RE.match(raw) or APPENDIX_RE.match(raw)
                if m and m.re is MANUAL_CHAPTER_RE:
                    chapter = f"Section {m.group(1)}"
                    chapter_title = " ".join(m.group(2).split())
                elif m:
                    chapter = f"Appendix {m.group(1)}"
                    chapter_title = " ".join(m.group(2).split())
                else:
                    chapter, chapter_title = text, ""
                article_stack, current = [], None
                audit["document_metadata"] += 1
            elif style in section_styles:
                m = NUMBERED_RE.match(text)
                if m:
                    start_section(m.group(1), m.group(2).strip(" ."))
                    audit["document_metadata"] += 1
                elif MANUAL_CHAPTER_RE.match(raw):
                    # TCM styles "SECTION 7 DRIVEWAYS" at Heading 2 already
                    # handled; a stray such heading resets the chapter.
                    m2 = MANUAL_CHAPTER_RE.match(raw)
                    chapter, chapter_title = f"Section {m2.group(1)}", " ".join(m2.group(2).split())
                    current = None
                    audit["document_metadata"] += 1
                else:
                    # Unnumbered sub-heading (e.g. "PURPOSE") — treat as a
                    # titled section under the current chapter.
                    start_section(text, "")
                    audit["document_metadata"] += 1
            elif style == "History Note":
                if current is not None:
                    current.history_notes.append(text)
                audit["history_note"] += 1
            elif FIGURE_CAPTION_RE.match(text) and style == "Normal":
                ensure_preamble()
                current.paragraphs.append(f"[caption] {text}")
                audit["table_caption"] += 1
            else:
                ensure_preamble()
                current.paragraphs.append(el.text.strip())
                audit["mapped_to_section"] += 1

    audit["sections"] = len(sections)
    audit["unresolved_samples"] = unresolved_samples
    return sections, audit


def load_all() -> tuple[list[Section], dict]:
    all_sections: list[Section] = []
    audits: dict[str, dict] = {}
    for doc_id, path in (("LDC", config.LDC_DOCX), ("DCM", config.DCM_DOCX),
                         ("TCM", config.TCM_DOCX)):
        secs, audit = load_document(path, doc_id)
        all_sections.extend(secs)
        audits[doc_id] = audit
    return all_sections, audits

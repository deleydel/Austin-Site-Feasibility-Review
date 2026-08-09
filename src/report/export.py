"""Export a report document to DOCX, HTML, or PDF."""

from __future__ import annotations

import html
import re
from pathlib import Path
from typing import Any, Literal

from docx import Document
from docx.shared import Inches, Pt
from fpdf import FPDF

from src.report.schema import build_report_document
from src.report.template import render_report_markdown

ExportFormat = Literal["docx", "html", "pdf", "md"]


def export_report(
    final_report_or_document: dict[str, Any],
    output_path: str | Path,
    fmt: ExportFormat | None = None,
) -> Path:
    """Export a guarded final_report (or built document) to disk.

    Format is inferred from the file suffix when ``fmt`` is omitted.
    """

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    if fmt is None:
        suffix = path.suffix.lower().lstrip(".")
        if suffix not in {"docx", "html", "pdf", "md"}:
            raise ValueError(
                f"Unsupported export suffix '{path.suffix}'. "
                "Use .docx, .html, .pdf, or .md"
            )
        fmt = suffix  # type: ignore[assignment]

    document = (
        final_report_or_document
        if "sections" in final_report_or_document
        else build_report_document(final_report_or_document)
    )

    if fmt == "md":
        path.write_text(render_report_markdown(document), encoding="utf-8")
    elif fmt == "html":
        path.write_text(render_report_html(document), encoding="utf-8")
    elif fmt == "docx":
        _write_docx(document, path)
    elif fmt == "pdf":
        _write_pdf(document, path)
    else:
        raise ValueError(f"Unsupported export format: {fmt}")

    return path


def render_report_html(document: dict[str, Any]) -> str:
    """Render the report document as a simple standalone HTML page."""

    md = render_report_markdown(document)
    body = _markdownish_to_html(md)
    title = html.escape(str(document.get("title") or "Feasibility Report"))
    return (
        "<!DOCTYPE html>\n"
        '<html lang="en">\n'
        "<head>\n"
        '  <meta charset="utf-8" />\n'
        f"  <title>{title}</title>\n"
        "  <style>\n"
        "    body { font-family: Georgia, 'Times New Roman', serif; "
        "margin: 2rem auto; max-width: 820px; line-height: 1.5; "
        "color: #1a1a1a; background: #faf8f5; padding: 0 1.25rem 3rem; }\n"
        "    h1, h2, h3 { font-family: 'Helvetica Neue', Helvetica, Arial, "
        "sans-serif; color: #16324f; }\n"
        "    h1 { border-bottom: 3px solid #c45c26; padding-bottom: 0.4rem; }\n"
        "    h2 { margin-top: 1.8rem; border-bottom: 1px solid #d8d2c8; "
        "padding-bottom: 0.25rem; }\n"
        "    ul { padding-left: 1.2rem; }\n"
        "    .disclaimer { background: #fff4e8; border-left: 4px solid "
        "#c45c26; padding: 0.75rem 1rem; margin-top: 1.5rem; }\n"
        "  </style>\n"
        "</head>\n"
        f"<body>\n{body}\n</body>\n</html>\n"
    )


def _write_docx(document: dict[str, Any], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_heading(
        str(document.get("title") or "Preliminary Development Feasibility Report"),
        level=0,
    )
    for run in title.runs:
        run.font.size = Pt(18)

    doc.add_paragraph(
        f"Generated: {document.get('generated_date', '')} | "
        f"Status: {document.get('status', '')}"
    )

    def add_bold_runs(paragraph, text: str) -> None:
        """Add text to a paragraph, rendering **bold** markdown as real bold runs."""

        for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
            if not part:
                continue
            run = paragraph.add_run(part)
            run.bold = bool(i % 2)

    md = render_report_markdown(document)
    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            continue  # already added as document title
        if line.startswith("Generated:") or line.startswith("Status:"):
            continue  # already added above
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_bold_runs(p, line[2:].strip())
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_bold_runs(p, re.sub(r"^\d+\.\s*", "", line))
        elif "**" in line:
            p = doc.add_paragraph()
            add_bold_runs(p, line)
        else:
            doc.add_paragraph(line)

    doc.save(path)


def _write_pdf(document: dict[str, Any], path: Path) -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    usable = pdf.w - pdf.l_margin - pdf.r_margin

    def write_line(text: str, height: float = 5) -> None:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(usable, height, _pdf_safe(text), align="L")

    def write_bold_runs(text: str, height: float = 5, indent: float = 0) -> None:
        """Render a line with **bold** markdown segments as real bold runs."""

        pdf.set_x(pdf.l_margin + indent)
        parts = re.split(r"\*\*(.+?)\*\*", _pdf_safe(text))
        base_size = pdf.font_size_pt
        for i, part in enumerate(parts):
            if not part:
                continue
            pdf.set_font("Helvetica", "B" if i % 2 else "", base_size)
            pdf.write(height, part)
        pdf.set_font("Helvetica", size=base_size)
        pdf.ln(height)

    pdf.set_font("Helvetica", "B", 16)
    title = str(
        document.get("title") or "Preliminary Development Feasibility Report"
    )
    write_line(title, 8)
    pdf.ln(2)
    pdf.set_font("Helvetica", size=11)
    write_line(
        f"Generated: {document.get('generated_date', '')} | "
        f"Status: {document.get('status', '')}",
        6,
    )
    pdf.ln(3)

    md = render_report_markdown(document)
    for raw_line in md.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("# "):
            continue
        if line.startswith("Generated:") or line.startswith("Status:"):
            continue
        if line.startswith("## "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 13)
            write_line(line[3:].strip(), 7)
            pdf.set_font("Helvetica", size=11)
        elif line.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            write_line(line[4:].strip(), 6)
            pdf.set_font("Helvetica", size=11)
        elif "**" in line:
            indent = 4 if line.startswith("- ") else 0
            write_bold_runs(line, 5, indent=indent)
        else:
            write_line(line, 5)

    pdf.output(str(path))


def _pdf_safe(text: str) -> str:
    """FPDF core fonts are Latin-1; replace unsupported characters."""

    return (
        text.replace("§", "Section ")
        .replace("—", "-")
        .replace("–", "-")
        .replace("“", '"')
        .replace("”", '"')
        .replace("’", "'")
        .encode("latin-1", errors="replace")
        .decode("latin-1")
    )


def _markdownish_to_html(md: str) -> str:
    """Minimal Markdown subset converter for report export."""

    parts: list[str] = []
    in_ul = False
    in_ol = False

    def close_lists() -> None:
        nonlocal in_ul, in_ol
        if in_ul:
            parts.append("</ul>")
            in_ul = False
        if in_ol:
            parts.append("</ol>")
            in_ol = False

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line:
            close_lists()
            continue
        if line.startswith("# "):
            close_lists()
            parts.append(f"<h1>{html.escape(line[2:])}</h1>")
        elif line.startswith("## "):
            close_lists()
            heading = html.escape(line[3:])
            if "disclaimer" in line.lower():
                parts.append(f'<h2>{heading}</h2><div class="disclaimer">')
            else:
                parts.append(f"<h2>{heading}</h2>")
        elif line.startswith("### "):
            close_lists()
            parts.append(f"<h3>{html.escape(line[4:])}</h3>")
        elif line.startswith("- "):
            if in_ol:
                parts.append("</ol>")
                in_ol = False
            if not in_ul:
                parts.append("<ul>")
                in_ul = True
            parts.append(f"<li>{_inline_md(line[2:])}</li>")
        elif re.match(r"^\d+\.\s", line):
            if in_ul:
                parts.append("</ul>")
                in_ul = False
            if not in_ol:
                parts.append("<ol>")
                in_ol = True
            text = re.sub(r"^\d+\.\s*", "", line)
            parts.append(f"<li>{_inline_md(text)}</li>")
        else:
            close_lists()
            parts.append(f"<p>{_inline_md(line)}</p>")

    close_lists()
    # Close disclaimer div if opened.
    html_body = "\n".join(parts)
    if 'class="disclaimer"' in html_body and not html_body.endswith("</div>"):
        html_body += "\n</div>"
    return html_body


def _inline_md(text: str) -> str:
    escaped = html.escape(text)
    return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
